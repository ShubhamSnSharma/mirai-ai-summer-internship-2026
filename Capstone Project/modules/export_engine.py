"""
Document Export Engine Orchestrator & Caching Pipeline.

Purpose:
    Serves strictly as an orchestration and caching layer for PDF and DOCX document generation.
    Delegates layout-specific rendering to template modules registered in templates/.
"""

import html
import io
import logging
import re
from datetime import datetime, timezone
from typing import Dict, Any, Optional, Tuple
import streamlit as st

# ReportLab Imports
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, HRFlowable
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_LEFT

# Python-Docx Imports
from docx import Document
from docx.shared import Inches

# Import Centralized Template Registry
from templates import get_template_module

# Initialize Logger
logger = logging.getLogger("export_engine")


def generate_filename(candidate_name: str, template_name: str, ext: str) -> str:
    """Generates clean, standardized export filenames.

    Args:
        candidate_name: Candidate full name string.
        template_name: Selected template identifier or display name.
        ext: Target file extension ('pdf' or 'docx').

    Returns:
        Formatted filename string (e.g. Alex_Chen_Modern_Professional.pdf).
    """
    clean_c = re.sub(r"[^\w\s-]", "", candidate_name or "Candidate").strip().replace(" ", "_")
    clean_t = re.sub(r"[^\w\s-]", "", template_name or "Template").strip().replace(" ", "_")
    ext_clean = ext.strip().lstrip(".")
    return f"{clean_c}_{clean_t}.{ext_clean}"


def build_pdf_summary_page(analysis_data: Dict[str, Any], styles: Any, options: Dict[str, bool]) -> list:
    """Builds flowables for optional AI Executive Summary cover page in PDF with html.escape XML safety.

    Args:
        analysis_data: Analysis dictionary conforming to json_contract.md.
        styles: ReportLab stylesheet object.
        options: Dictionary of checkboxes (include_summary, include_verdict, include_scores).

    Returns:
        List of ReportLab Flowable objects.
    """
    story = []
    candidate = analysis_data.get("candidate", {})
    job = analysis_data.get("job", {})
    scores = analysis_data.get("scores", {})
    recruiter = analysis_data.get("recruiter_feedback", {})
    metadata = analysis_data.get("metadata", {})

    c_name = html.escape(candidate.get("name", "Candidate"))
    j_role = html.escape(job.get("role", "Target Role"))
    j_company = html.escape(job.get("company", ""))
    model_name = html.escape(metadata.get("model", "gemini-2.5-flash"))

    title_style = ParagraphStyle(
        "ReportTitle",
        parent=styles["Heading1"],
        fontName="Helvetica-Bold",
        fontSize=20,
        leading=24,
        textColor=colors.HexColor("#1E293B"),
        alignment=TA_LEFT,
        spaceAfter=6,
    )
    
    story.append(Paragraph("AI Resume Analysis Executive Report", title_style))
    story.append(Paragraph(f"Candidate: <b>{c_name}</b> | Target Role: <b>{j_role} ({j_company})</b>", styles["Normal"]))
    raw_timestamp = metadata.get('analysis_timestamp')
    if raw_timestamp and isinstance(raw_timestamp, str) and len(raw_timestamp) >= 10:
        analysis_date = raw_timestamp[:10]
    else:
        analysis_date = datetime.now().strftime("%Y-%m-%d")

    story.append(Paragraph(f"Analysis Date: {analysis_date} | Engine: {model_name}", styles["Normal"]))
    story.append(Spacer(1, 12))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#CBD5E1"), spaceAfter=15))

    if options.get("include_scores", True):
        story.append(Paragraph("<b>Quantitative Evaluation Scores</b>", styles["Heading2"]))
        score_data = [
            ["Metric", "Score", "Rating Benchmark"],
            ["Overall Resume Score", f"{scores.get('overall_resume_score', 0)} / 100", "Weighted Composite"],
            ["ATS Compliance Score", f"{scores.get('ats_score', 0)} / 100", "Parser Pass Rate"],
            ["Job Match Score", f"{scores.get('job_match_score', 0)} / 100", "Skill & Keyword Fit"],
            ["Interview Callback Probability", f"{scores.get('interview_probability', 0)}%", "Probability Estimate"],
        ]
        t = Table(score_data, colWidths=[200, 100, 180])
        t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#1E293B")),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 6),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor("#E2E8F0")),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
        ]))
        story.append(t)
        story.append(Spacer(1, 15))

    if options.get("include_verdict", True):
        story.append(Paragraph("<b>Recruiter Verdict & Recommendation</b>", styles["Heading2"]))
        story.append(Paragraph(f"<b>Verdict:</b> {html.escape(recruiter.get('overall_verdict', ''))}", styles["Normal"]))
        story.append(Paragraph(f"<b>Hiring Decision:</b> {html.escape(recruiter.get('hire_decision', ''))}", styles["Normal"]))
        story.append(Spacer(1, 6))
        story.append(Paragraph(f"<b>Recruiter Commentary:</b> <i>{html.escape(recruiter.get('final_comments', ''))}</i>", styles["Normal"]))
        story.append(Spacer(1, 15))

    if options.get("include_summary", True):
        story.append(Paragraph("<b>Top Strengths & Key Concerns</b>", styles["Heading2"]))
        strengths = analysis_data.get("strengths", [])
        weaknesses = analysis_data.get("weaknesses", [])
        
        story.append(Paragraph("<b>Strengths:</b>", styles["Normal"]))
        for s in strengths[:3]:
            story.append(Paragraph(f"• {html.escape(s)}", styles["Normal"]))
        story.append(Spacer(1, 6))
        story.append(Paragraph("<b>Areas for Improvement:</b>", styles["Normal"]))
        for w in weaknesses[:3]:
            story.append(Paragraph(f"• {html.escape(w)}", styles["Normal"]))

    story.append(PageBreak())
    return story


def generate_pdf(analysis_data: Dict[str, Any], template_name: str = "modern_professional", options: Optional[Dict[str, bool]] = None) -> bytes:
    """Orchestrates PDF document generation by delegating layout rendering to the selected template module.

    Args:
        analysis_data: Analysis dictionary conforming to json_contract.md.
        template_name: Selected template name ('ats_professional', 'modern_professional', 'developer_professional').
        options: Optional export checkboxes dictionary.

    Returns:
        Binary bytes of the generated PDF document.
    """
    if options is None:
        options = {}

    opt_resume = analysis_data.get("optimized_resume", {})
    template_mod = get_template_module(template_name)
    
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=36,
        leftMargin=36,
        topMargin=36,
        bottomMargin=36,
    )

    styles = getSampleStyleSheet()
    story = []

    # Prepend Optional Summary Cover Page if requested
    if any(options.values()):
        summary_story = build_pdf_summary_page(analysis_data, styles, options)
        story.extend(summary_story)

    # Delegate Layout Rendering to Selected Template Module
    template_mod.render_pdf(opt_resume, styles, story)

    # Build PDF Document
    try:
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()
    except Exception as e:
        logger.exception("ReportLab PDF generation failed for template '%s': %s", template_name, str(e))
        raise RuntimeError(f"ReportLab PDF generation error: {str(e)}")


def generate_docx(analysis_data: Dict[str, Any], template_name: str = "modern_professional", options: Optional[Dict[str, bool]] = None) -> bytes:
    """Orchestrates DOCX document generation by delegating layout rendering to the selected template module.

    Args:
        analysis_data: Analysis dictionary conforming to json_contract.md.
        template_name: Selected template name.
        options: Optional export checkboxes dictionary.

    Returns:
        Binary bytes of the generated DOCX document.
    """
    if options is None:
        options = {}

    opt_resume = analysis_data.get("optimized_resume", {})
    template_mod = get_template_module(template_name)

    doc = Document()

    # Set standard 0.75-inch document margins
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Prepend Optional Summary Cover Page if requested
    if any(options.values()):
        candidate = analysis_data.get("candidate", {})
        job = analysis_data.get("job", {})
        scores = analysis_data.get("scores", {})
        recruiter = analysis_data.get("recruiter_feedback", {})

        doc.add_heading("AI Resume Analysis Executive Report", level=1)
        doc.add_paragraph(f"Candidate: {candidate.get('name', '')} | Target Role: {job.get('role', '')} ({job.get('company', '')})")
        
        if options.get("include_scores", True):
            doc.add_heading("Evaluation Scores", level=2)
            p_scores = doc.add_paragraph()
            p_scores.add_run(f"Overall Score: {scores.get('overall_resume_score', 0)} / 100\n").bold = True
            p_scores.add_run(f"ATS Score: {scores.get('ats_score', 0)} / 100\n")
            p_scores.add_run(f"Job Match Score: {scores.get('job_match_score', 0)} / 100\n")

        if options.get("include_verdict", True):
            doc.add_heading("Recruiter Verdict", level=2)
            doc.add_paragraph(f"Verdict: {recruiter.get('overall_verdict', '')}")
            doc.add_paragraph(f"Hiring Decision: {recruiter.get('hire_decision', '')}")
            doc.add_paragraph(f"Comments: {recruiter.get('final_comments', '')}")

        doc.add_page_break()

    # Delegate Layout Rendering to Selected Template Module
    template_mod.render_docx(opt_resume, doc)

    # Build DOCX Buffer
    buffer = io.BytesIO()
    try:
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    except Exception as e:
        logger.exception("Python-Docx DOCX generation failed for template '%s': %s", template_name, str(e))
        raise RuntimeError(f"Python-Docx DOCX generation error: {str(e)}")
