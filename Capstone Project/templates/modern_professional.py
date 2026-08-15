"""
Modern Professional Resume Template Renderer & Metadata Specification.

Single Source of Truth for Modern Professional Executive 2-Column Layout.
Exposes metadata properties and rendering methods for Streamlit Live Preview, ReportLab PDF, and Python-Docx DOCX.
"""

import html
from typing import Dict, Any
import streamlit as st

# ReportLab Imports
from reportlab.platypus import Paragraph, Spacer, HRFlowable, Table, TableStyle
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib import colors

# Python-Docx Imports
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Module Metadata Specification
KEY: str = "modern_professional"
DISPLAY_NAME: str = "Modern Professional"
DESCRIPTION: str = "Executive 2-column corporate layout with subtle accent dividers."
ATS_COMPATIBILITY: str = "95% ATS Safe"
RECOMMENDED_FOR: str = "Senior, Product & Executive roles"
ESTIMATED_PAGES: str = "1 Page"
EXPORT_READY: bool = True
PREVIEW_THUMBNAIL: str = "┌── 2-Column Split ─┐\n│ Skills │ Experi   │\n│ Edu    │ Projects │\n└────────┴──────────┘"


def render_streamlit(resume_data: Dict[str, Any]) -> None:
    """Renders 2-column Modern Professional executive resume layout for Streamlit preview.

    Args:
        resume_data: Dictionary containing candidate optimized_resume structure.
    """
    personal = resume_data.get("personal_information", {})
    c_name = personal.get("name", "Candidate Name")
    headline = resume_data.get("headline", "Professional Title")
    email = personal.get("email", "")
    phone = personal.get("phone", "")
    location = personal.get("location", "")
    linkedin = personal.get("linkedin", "")
    github = personal.get("github", "")
    portfolio = personal.get("portfolio", "")

    # Executive Banner
    st.markdown(f"# {c_name}")
    st.markdown(f"### {headline}")
    st.caption(f"📍 {location}  •  ✉️ {email}  •  📞 {phone}  •  🔗 {linkedin}  •  🌐 {portfolio}")
    st.divider()

    # 2-Column Layout
    col_left, col_right = st.columns([1, 2])

    with col_left:
        summary = resume_data.get("professional_summary", "")
        if summary:
            st.markdown("#### Executive Summary")
            st.write(summary)
            st.divider()

        skills = resume_data.get("skills", {})
        if skills:
            st.markdown("#### Core Competencies")
            for cat, items in skills.items():
                if items:
                    cat_name = cat.replace("_", " ").title()
                    st.markdown(f"**{cat_name}:**")
                    st.write(", ".join(items))
            st.divider()

        education = resume_data.get("education", [])
        if education:
            st.markdown("#### Education")
            for edu in education:
                degree = edu.get("degree", "")
                inst = edu.get("institution", "")
                grad = edu.get("graduation_date", "")
                gpa = edu.get("gpa", "")

                st.markdown(f"**{degree}**")
                st.write(inst)
                st.caption(f"Graduated: {grad} • GPA: {gpa}")
            st.divider()

        certs = resume_data.get("certifications", [])
        achievements = resume_data.get("achievements", [])
        if certs or achievements:
            if certs:
                st.markdown("#### Certifications")
                for c in certs:
                    st.markdown(f"- 🏆 {c}")
            if achievements:
                st.markdown("#### Key Achievements")
                for a in achievements:
                    st.markdown(f"- ⭐ {a}")

    with col_right:
        experiences = resume_data.get("experience", [])
        if experiences:
            st.markdown("#### Professional Experience")
            for exp in experiences:
                company = exp.get("company", "")
                role = exp.get("role", "")
                start = exp.get("start_date", "")
                end = exp.get("end_date", "")
                loc = exp.get("location", "")

                st.markdown(f"##### {role}")
                st.markdown(f"**{company}** | *{loc}*")
                st.caption(f"📅 {start} – {end}")

                for b in exp.get("bullets", []):
                    st.markdown(f"- {b}")
                st.write("")
            st.divider()

        projects = resume_data.get("projects", [])
        if projects:
            st.markdown("#### High-Impact Projects")
            for proj in projects:
                p_name = proj.get("project_name", "")
                tech_stack = ", ".join(proj.get("tech_stack", []))
                desc = proj.get("description", "")

                st.markdown(f"##### 📌 {p_name}")
                if tech_stack:
                    st.caption(f"**Tech Stack:** `{tech_stack}`")
                st.write(desc)
                st.write("")


def render_pdf(resume_data: Dict[str, Any], styles: Any, story: list) -> None:
    """Appends Modern Professional executive 2-column ReportLab flowables to story.
    Applies html.escape() to protect against XML special character crashes.

    Args:
        resume_data: Candidate optimized_resume structure.
        styles: ReportLab sample stylesheet.
        story: List of flowable elements to append to.
    """
    personal = resume_data.get("personal_information", {})
    c_name = html.escape(personal.get("name", "Candidate Name"))
    headline = html.escape(resume_data.get("headline", "Professional Title"))
    email = html.escape(personal.get("email", ""))
    phone = html.escape(personal.get("phone", ""))
    location = html.escape(personal.get("location", ""))
    linkedin = html.escape(personal.get("linkedin", ""))
    github = html.escape(personal.get("github", ""))

    accent_blue = colors.HexColor("#2563EB")
    dark_slate = colors.HexColor("#1E293B")

    title_style = ParagraphStyle("Mod_Title", parent=styles["Heading1"], fontName="Helvetica-Bold", fontSize=20, leading=24, textColor=accent_blue, spaceAfter=2)
    sub_style = ParagraphStyle("Mod_Sub", parent=styles["Normal"], fontName="Helvetica-Bold", fontSize=11, leading=14, textColor=dark_slate, spaceAfter=4)
    contact_style = ParagraphStyle("Mod_Contact", parent=styles["Normal"], fontName="Helvetica", fontSize=9, leading=12, textColor=colors.HexColor("#64748B"), spaceAfter=6)
    section_h2 = ParagraphStyle("Mod_H2", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=11, leading=14, textColor=accent_blue, spaceBefore=6, spaceAfter=3)
    body_style = ParagraphStyle("Mod_Body", parent=styles["Normal"], fontName="Helvetica", fontSize=9, leading=12, textColor=dark_slate, spaceAfter=3)
    bullet_style = ParagraphStyle("Mod_Bullet", parent=styles["Normal"], fontName="Helvetica", fontSize=8.5, leading=11, textColor=colors.HexColor("#334155"), leftIndent=8, spaceAfter=2)

    # Top Banner Header
    story.append(Paragraph(c_name, title_style))
    story.append(Paragraph(headline, sub_style))
    story.append(Paragraph(f"{location} • {email} • {phone} • {linkedin} • {github}", contact_style))
    story.append(HRFlowable(width="100%", thickness=1.5, color=accent_blue, spaceAfter=8))

    # Left Column Flowables
    left_flowables = []
    summary = resume_data.get("professional_summary", "")
    if summary:
        left_flowables.append(Paragraph("EXECUTIVE SUMMARY", section_h2))
        left_flowables.append(Paragraph(html.escape(summary), body_style))
        left_flowables.append(Spacer(1, 4))

    skills = resume_data.get("skills", {})
    if skills:
        left_flowables.append(Paragraph("CORE COMPETENCIES", section_h2))
        for cat, items in skills.items():
            if items:
                escaped_cat = html.escape(cat.replace("_", " ").title())
                escaped_items = ", ".join([html.escape(i) for i in items])
                left_flowables.append(Paragraph(f"<b>{escaped_cat}:</b>", body_style))
                left_flowables.append(Paragraph(escaped_items, body_style))
        left_flowables.append(Spacer(1, 4))

    education = resume_data.get("education", [])
    if education:
        left_flowables.append(Paragraph("EDUCATION", section_h2))
        for edu in education:
            deg = html.escape(edu.get("degree", ""))
            inst = html.escape(edu.get("institution", ""))
            grad = html.escape(edu.get("graduation_date", ""))
            left_flowables.append(Paragraph(f"<b>{deg}</b>", body_style))
            left_flowables.append(Paragraph(f"{inst} ({grad})", body_style))

    # Right Column Flowables
    right_flowables = []
    experiences = resume_data.get("experience", [])
    if experiences:
        right_flowables.append(Paragraph("PROFESSIONAL EXPERIENCE", section_h2))
        for exp in experiences:
            comp = html.escape(exp.get("company", ""))
            role = html.escape(exp.get("role", ""))
            start = html.escape(exp.get("start_date", ""))
            end = html.escape(exp.get("end_date", ""))
            loc = html.escape(exp.get("location", ""))

            right_flowables.append(Paragraph(f"<b>{role}</b> — {comp}", body_style))
            right_flowables.append(Paragraph(f"<i>{loc} | {start} – {end}</i>", contact_style))

            for b in exp.get("bullets", []):
                right_flowables.append(Paragraph(f"• {html.escape(b)}", bullet_style))
            right_flowables.append(Spacer(1, 4))

    projects = resume_data.get("projects", [])
    if projects:
        right_flowables.append(Paragraph("HIGH-IMPACT PROJECTS", section_h2))
        for proj in projects:
            p_name = html.escape(proj.get("project_name", ""))
            tech_stack = ", ".join([html.escape(t) for t in proj.get("tech_stack", [])])
            desc = html.escape(proj.get("description", ""))

            right_flowables.append(Paragraph(f"<b>{p_name}</b> {f'[{tech_stack}]' if tech_stack else ''}", body_style))
            right_flowables.append(Paragraph(desc, bullet_style))
            right_flowables.append(Spacer(1, 3))

    # Package 2 Columns into Table
    grid_table = Table([[left_flowables, right_flowables]], colWidths=[175, 325])
    grid_table.setStyle(TableStyle([
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('LEFTPADDING', (0, 0), (-1, -1), 0),
        ('RIGHTPADDING', (0, 0), (-1, -1), 6),
        ('TOPPADDING', (0, 0), (-1, -1), 0),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
    ]))

    story.append(grid_table)


def render_docx(resume_data: Dict[str, Any], doc: Any) -> None:
    """Appends Modern Professional executive 2-column layout to Python-Docx Document.

    Args:
        resume_data: Candidate optimized_resume structure.
        doc: Python-Docx Document instance.
    """
    personal = resume_data.get("personal_information", {})
    c_name = personal.get("name", "Candidate Name")
    headline = resume_data.get("headline", "Professional Title")
    email = personal.get("email", "")
    phone = personal.get("phone", "")
    location = personal.get("location", "")
    linkedin = personal.get("linkedin", "")
    github = personal.get("github", "")

    # Top Executive Banner
    h_name = doc.add_heading(c_name, level=0)
    h_name.alignment = WD_ALIGN_PARAGRAPH.LEFT

    p_head = doc.add_paragraph()
    r_head = p_head.add_run(headline)
    r_head.bold = True
    r_head.font.size = Pt(12)
    r_head.font.color.rgb = RGBColor(37, 99, 235)

    p_contact = doc.add_paragraph(f"{location} • {email} • {phone} • {linkedin} • {github}")
    p_contact.style.font.size = Pt(9.5)
    p_contact.style.font.color.rgb = RGBColor(100, 116, 139)

    # 2-Column Table Split
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False

    cell_left = table.rows[0].cells[0]
    cell_right = table.rows[0].cells[1]

    cell_left.width = Inches(2.3)
    cell_right.width = Inches(4.7)

    # Left Column: Summary, Skills, Education
    p_sum_h = cell_left.add_paragraph()
    r_sum_h = p_sum_h.add_run("EXECUTIVE SUMMARY")
    r_sum_h.bold = True
    r_sum_h.font.color.rgb = RGBColor(37, 99, 235)

    summary = resume_data.get("professional_summary", "")
    if summary:
        cell_left.add_paragraph(summary)

    p_sk_h = cell_left.add_paragraph()
    r_sk_h = p_sk_h.add_run("CORE COMPETENCIES")
    r_sk_h.bold = True
    r_sk_h.font.color.rgb = RGBColor(37, 99, 235)

    skills = resume_data.get("skills", {})
    if skills:
        for cat, items in skills.items():
            if items:
                p_sk = cell_left.add_paragraph()
                p_sk.add_run(f"{cat.replace('_', ' ').title()}:\n").bold = True
                p_sk.add_run(", ".join(items))

    education = resume_data.get("education", [])
    if education:
        p_edu_h = cell_left.add_paragraph()
        r_edu_h = p_edu_h.add_run("EDUCATION")
        r_edu_h.bold = True
        r_edu_h.font.color.rgb = RGBColor(37, 99, 235)

        for edu in education:
            deg = edu.get("degree", "")
            inst = edu.get("institution", "")
            grad = edu.get("graduation_date", "")
            cell_left.add_paragraph(f"{deg}\n{inst} ({grad})")

    # Right Column: Experience, Projects
    p_exp_h = cell_right.add_paragraph()
    r_exp_h = p_exp_h.add_run("PROFESSIONAL EXPERIENCE")
    r_exp_h.bold = True
    r_exp_h.font.color.rgb = RGBColor(37, 99, 235)

    experiences = resume_data.get("experience", [])
    if experiences:
        for exp in experiences:
            comp = exp.get("company", "")
            role = exp.get("role", "")
            start = exp.get("start_date", "")
            end = exp.get("end_date", "")
            loc = exp.get("location", "")

            p_exp = cell_right.add_paragraph()
            r_role = p_exp.add_run(f"{role} — {comp}")
            r_role.bold = True
            p_exp.add_run(f"\n{loc} | {start} – {end}").italic = True

            for b in exp.get("bullets", []):
                cell_right.add_paragraph(b, style="List Bullet")

    projects = resume_data.get("projects", [])
    if projects:
        p_proj_h = cell_right.add_paragraph()
        r_proj_h = p_proj_h.add_run("HIGH-IMPACT PROJECTS")
        r_proj_h.bold = True
        r_proj_h.font.color.rgb = RGBColor(37, 99, 235)

        for proj in projects:
            p_name = proj.get("project_name", "")
            tech_stack = ", ".join(proj.get("tech_stack", []))
            desc = proj.get("description", "")

            p_proj = cell_right.add_paragraph()
            p_proj.add_run(p_name).bold = True
            if tech_stack:
                p_proj.add_run(f" [{tech_stack}]").italic = True
            cell_right.add_paragraph(desc)
