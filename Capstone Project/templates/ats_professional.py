"""
ATS Professional Resume Template Renderer & Metadata Specification.

Single Source of Truth for ATS Professional Layout.
Exposes metadata properties and rendering methods for Streamlit Live Preview, ReportLab PDF, and Python-Docx DOCX.
"""

import html
from typing import Dict, Any
import streamlit as st

# ReportLab Imports
from reportlab.platypus import Paragraph, Spacer, HRFlowable
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib import colors

# Python-Docx Imports
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Module Metadata Specification
KEY: str = "ats_professional"
DISPLAY_NAME: str = "ATS Professional"
DESCRIPTION: str = "Single-column, high-parseability layout designed for 100% ATS pass rates."
ATS_COMPATIBILITY: str = "100% ATS Safe"
RECOMMENDED_FOR: str = "Corporate & High-volume applications"
ESTIMATED_PAGES: str = "1 Page"
EXPORT_READY: bool = True
PREVIEW_THUMBNAIL: str = "┌── Single Column ──┐\n│ Header Info      │\n│ Summary & Skills │\n│ Experience       │\n└──────────────────┘"


def render_streamlit(resume_data: Dict[str, Any]) -> None:
    """Renders single-column ATS-friendly resume layout for Streamlit preview.

    Args:
        resume_data: Dictionary containing candidate optimized_resume structure.
    """
    personal = resume_data.get("personal_information", {})
    c_name = personal.get("name", "Candidate Name")
    headline = resume_data.get("headline", "Professional Title")
    email = personal.get("email", "")
    phone = personal.get("phone", "")
    location = personal.get("location", "")
    linkedin = personal.get("linkedin", "")
    github = personal.get("github", "")

    # Header / Contact Section
    st.markdown(f"## {c_name}")
    st.markdown(f"**{headline}**")
    st.caption(f"{location} | {email} | {phone} | {linkedin} | {github}")
    st.divider()

    # Professional Summary
    summary = resume_data.get("professional_summary", "")
    if summary:
        st.markdown("### Professional Summary")
        st.write(summary)
        st.divider()

    # Technical Skills Matrix
    skills = resume_data.get("skills", {})
    if skills:
        st.markdown("### Technical Skills")
        for category, items in skills.items():
            if items:
                cat_title = category.replace("_", " ").title()
                st.markdown(f"**{cat_title}:** {', '.join(items)}")
        st.divider()

    # Experience Section
    experiences = resume_data.get("experience", [])
    if experiences:
        st.markdown("### Professional Experience")
        for exp in experiences:
            company = exp.get("company", "")
            role = exp.get("role", "")
            start = exp.get("start_date", "")
            end = exp.get("end_date", "")
            loc = exp.get("location", "")

            st.markdown(f"#### {role} — {company}")
            st.caption(f"{loc} | {start} – {end}")

            for b in exp.get("bullets", []):
                st.markdown(f"- {b}")
            st.write("")
        st.divider()

    # Technical Projects Section
    projects = resume_data.get("projects", [])
    if projects:
        st.markdown("### Technical Projects")
        for proj in projects:
            p_name = proj.get("project_name", "")
            tech_stack = ", ".join(proj.get("tech_stack", []))
            desc = proj.get("description", "")

            st.markdown(f"#### {p_name}")
            if tech_stack:
                st.caption(f"**Technologies:** {tech_stack}")
            st.write(desc)
            st.write("")
        st.divider()

    # Education Section
    education = resume_data.get("education", [])
    if education:
        st.markdown("### Education")
        for edu in education:
            degree = edu.get("degree", "")
            inst = edu.get("institution", "")
            grad = edu.get("graduation_date", "")
            gpa = edu.get("gpa", "")

            st.markdown(f"**{degree}** — {inst}")
            st.caption(f"Graduated: {grad} | GPA: {gpa}")
        st.divider()

    # Certifications & Achievements
    certs = resume_data.get("certifications", [])
    achievements = resume_data.get("achievements", [])
    if certs or achievements:
        if certs:
            st.markdown("### Certifications")
            for c in certs:
                st.markdown(f"- {c}")
        if achievements:
            st.markdown("### Achievements")
            for a in achievements:
                st.markdown(f"- {a}")


def render_pdf(resume_data: Dict[str, Any], styles: Any, story: list) -> None:
    """Appends ATS Professional single-column ReportLab flowables to story.
    Applies html.escape() to protect against XML special character crashes.

    Args:
        resume_data: Candidate optimized_resume structure.
        styles: ReportLab sample stylesheet.
        story: List of flowable elements to append to.
    """
    personal = resume_data.get("personal_information", {})
    c_name = html.escape(personal.get("name", "Candidate Name"))
    headline = html.escape(resume_data.get("headline", "Professional Title"))
    email = html.escape(personal.get("email", ""))
    phone = html.escape(personal.get("phone", ""))
    location = html.escape(personal.get("location", ""))
    linkedin = html.escape(personal.get("linkedin", ""))
    github = html.escape(personal.get("github", ""))

    title_style = ParagraphStyle(
        "ATS_Title",
        parent=styles["Heading1"],
        fontName="Helvetica-Bold",
        fontSize=18,
        leading=22,
        textColor=colors.HexColor("#0F172A"),
        spaceAfter=2,
    )
    sub_style = ParagraphStyle(
        "ATS_Sub",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=11,
        leading=14,
        textColor=colors.HexColor("#334155"),
        spaceAfter=4,
    )
    contact_style = ParagraphStyle(
        "ATS_Contact",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=9,
        leading=12,
        textColor=colors.HexColor("#475569"),
        spaceAfter=8,
    )
    heading_style = ParagraphStyle(
        "ATS_Heading",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=11,
        leading=15,
        textColor=colors.HexColor("#0F172A"),
        spaceBefore=8,
        spaceAfter=3,
    )
    body_style = ParagraphStyle(
        "ATS_Body",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=9.5,
        leading=13,
        textColor=colors.HexColor("#1E293B"),
        spaceAfter=3,
    )
    bullet_style = ParagraphStyle(
        "ATS_Bullet",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=9,
        leading=12,
        textColor=colors.HexColor("#334155"),
        leftIndent=10,
        spaceAfter=2,
    )

    story.append(Paragraph(c_name, title_style))
    story.append(Paragraph(headline, sub_style))
    story.append(Paragraph(f"{location} | {email} | {phone} | {linkedin} | {github}", contact_style))
    story.append(HRFlowable(width="100%", thickness=0.8, color=colors.HexColor("#94A3B8"), spaceAfter=8))

    # Summary
    summary = resume_data.get("professional_summary", "")
    if summary:
        story.append(Paragraph("PROFESSIONAL SUMMARY", heading_style))
        story.append(Paragraph(html.escape(summary), body_style))
        story.append(Spacer(1, 4))

    # Skills
    skills = resume_data.get("skills", {})
    if skills:
        story.append(Paragraph("TECHNICAL SKILLS", heading_style))
        for cat, items in skills.items():
            if items:
                escaped_cat = html.escape(cat.replace("_", " ").title())
                escaped_items = ", ".join([html.escape(i) for i in items])
                story.append(Paragraph(f"<b>{escaped_cat}:</b> {escaped_items}", body_style))
        story.append(Spacer(1, 4))

    # Experience
    experiences = resume_data.get("experience", [])
    if experiences:
        story.append(Paragraph("PROFESSIONAL EXPERIENCE", heading_style))
        for exp in experiences:
            comp = html.escape(exp.get("company", ""))
            role = html.escape(exp.get("role", ""))
            start = html.escape(exp.get("start_date", ""))
            end = html.escape(exp.get("end_date", ""))
            loc = html.escape(exp.get("location", ""))

            story.append(Paragraph(f"<b>{role}</b> — {comp} ({loc}) | <i>{start} – {end}</i>", body_style))
            for b in exp.get("bullets", []):
                story.append(Paragraph(f"• {html.escape(b)}", bullet_style))
            story.append(Spacer(1, 3))

    # Projects
    projects = resume_data.get("projects", [])
    if projects:
        story.append(Paragraph("TECHNICAL PROJECTS", heading_style))
        for proj in projects:
            p_name = html.escape(proj.get("project_name", ""))
            tech_stack = ", ".join([html.escape(t) for t in proj.get("tech_stack", [])])
            desc = html.escape(proj.get("description", ""))

            story.append(Paragraph(f"<b>{p_name}</b> {f'[{tech_stack}]' if tech_stack else ''}", body_style))
            story.append(Paragraph(desc, bullet_style))
            story.append(Spacer(1, 3))

    # Education
    education = resume_data.get("education", [])
    if education:
        story.append(Paragraph("EDUCATION", heading_style))
        for edu in education:
            deg = html.escape(edu.get("degree", ""))
            inst = html.escape(edu.get("institution", ""))
            grad = html.escape(edu.get("graduation_date", ""))
            gpa = html.escape(str(edu.get("gpa", "")))

            story.append(Paragraph(f"<b>{deg}</b> — {inst} ({grad}) {f'| GPA: {gpa}' if gpa else ''}", body_style))


def render_docx(resume_data: Dict[str, Any], doc: Any) -> None:
    """Appends ATS Professional single-column content to Python-Docx Document.

    Args:
        resume_data: Candidate optimized_resume structure.
        doc: Python-Docx Document instance.
    """
    personal = resume_data.get("personal_information", {})
    c_name = personal.get("name", "Candidate Name")
    headline = resume_data.get("headline", "Professional Title")
    email = personal.get("email", "")
    phone = personal.get("phone", "")
    location = personal.get("location", "")
    linkedin = personal.get("linkedin", "")
    github = personal.get("github", "")

    h_name = doc.add_heading(c_name, level=0)
    h_name.alignment = WD_ALIGN_PARAGRAPH.LEFT

    p_head = doc.add_paragraph()
    r_head = p_head.add_run(headline)
    r_head.bold = True
    r_head.font.size = Pt(11)

    p_contact = doc.add_paragraph(f"{location} | {email} | {phone} | {linkedin} | {github}")
    p_contact.style.font.size = Pt(9)
    p_contact.style.font.color.rgb = RGBColor(71, 85, 105)

    # Summary
    summary = resume_data.get("professional_summary", "")
    if summary:
        doc.add_heading("PROFESSIONAL SUMMARY", level=2)
        doc.add_paragraph(summary)

    # Skills
    skills = resume_data.get("skills", {})
    if skills:
        doc.add_heading("TECHNICAL SKILLS", level=2)
        for cat, items in skills.items():
            if items:
                p_sk = doc.add_paragraph()
                p_sk.add_run(f"{cat.replace('_', ' ').title()}: ").bold = True
                p_sk.add_run(", ".join(items))

    # Experience
    experiences = resume_data.get("experience", [])
    if experiences:
        doc.add_heading("PROFESSIONAL EXPERIENCE", level=2)
        for exp in experiences:
            comp = exp.get("company", "")
            role = exp.get("role", "")
            start = exp.get("start_date", "")
            end = exp.get("end_date", "")
            loc = exp.get("location", "")

            p_exp = doc.add_paragraph()
            r_role = p_exp.add_run(f"{role} — {comp}")
            r_role.bold = True
            p_exp.add_run(f" ({loc}) | {start} – {end}").italic = True

            for b in exp.get("bullets", []):
                doc.add_paragraph(b, style="List Bullet")

    # Projects
    projects = resume_data.get("projects", [])
    if projects:
        doc.add_heading("TECHNICAL PROJECTS", level=2)
        for proj in projects:
            p_name = proj.get("project_name", "")
            tech_stack = ", ".join(proj.get("tech_stack", []))
            desc = proj.get("description", "")

            p_proj = doc.add_paragraph()
            p_proj.add_run(p_name).bold = True
            if tech_stack:
                p_proj.add_run(f" [{tech_stack}]").italic = True
            doc.add_paragraph(desc)

    # Education
    education = resume_data.get("education", [])
    if education:
        doc.add_heading("EDUCATION", level=2)
        for edu in education:
            deg = edu.get("degree", "")
            inst = edu.get("institution", "")
            grad = edu.get("graduation_date", "")
            gpa = edu.get("gpa", "")

            p_edu = doc.add_paragraph()
            p_edu.add_run(f"{deg} — {inst}").bold = True
            p_edu.add_run(f" ({grad}) {f'GPA: {gpa}' if gpa else ''}")
