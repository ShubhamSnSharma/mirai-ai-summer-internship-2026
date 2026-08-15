"""
Developer Specialist Resume Template Renderer & Metadata Specification.

Single Source of Truth for Developer Specialist Tech-Focused Layout.
Exposes metadata properties and rendering methods for Streamlit Live Preview, ReportLab PDF, and Python-Docx DOCX.
"""

import html
from typing import Dict, Any
import streamlit as st

# ReportLab Imports
from reportlab.platypus import Paragraph, Spacer, HRFlowable
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib import colors

# Python-Docx Imports
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Module Metadata Specification
KEY: str = "developer_professional"
DISPLAY_NAME: str = "Developer Specialist"
DESCRIPTION: str = "Code-focused layout highlighting technical skills matrix & repo links."
ATS_COMPATIBILITY: str = "92% ATS Safe"
RECOMMENDED_FOR: str = "Software Engineers & AI Specialists"
ESTIMATED_PAGES: str = "1 Page"
EXPORT_READY: bool = True
PREVIEW_THUMBNAIL: str = "┌── Dev Repos ──────┐\n│ GitHub / Handles  │\n│ Tech Matrix Top   │\n│ Project Repos     │\n└───────────────────┘"


def render_streamlit(resume_data: Dict[str, Any]) -> None:
    """Renders Developer Specialist tech-focused resume layout for Streamlit preview.

    Args:
        resume_data: Dictionary containing candidate optimized_resume structure.
    """
    personal = resume_data.get("personal_information", {})
    c_name = personal.get("name", "Candidate Name")
    headline = resume_data.get("headline", "Full Stack AI Engineer")
    email = personal.get("email", "")
    phone = personal.get("phone", "")
    location = personal.get("location", "")
    linkedin = personal.get("linkedin", "")
    github = personal.get("github", "")
    portfolio = personal.get("portfolio", "")

    # Developer Header
    st.markdown(f"# 💻 {c_name}")
    st.markdown(f"**{headline}**")
    st.caption(f"📍 {location} | ✉️ {email} | 📞 {phone}")
    st.markdown(f"`GitHub: {github}` &nbsp;•&nbsp; `{portfolio}` &nbsp;•&nbsp; `{linkedin}`")
    st.divider()

    # Technical Skills Matrix at Top
    skills = resume_data.get("skills", {})
    if skills:
        st.markdown("### ⚙️ Developer Tech Matrix")
        col_dev1, col_dev2 = st.columns(2)
        with col_dev1:
            if skills.get("languages"):
                st.markdown(f"**Languages:** `{'`, `'.join(skills['languages'])}`")
            if skills.get("frameworks"):
                st.markdown(f"**Frameworks:** `{'`, `'.join(skills['frameworks'])}`")
            if skills.get("databases"):
                st.markdown(f"**Databases:** `{'`, `'.join(skills['databases'])}`")

        with col_dev2:
            if skills.get("tools"):
                st.markdown(f"**DevOps & Tools:** `{'`, `'.join(skills['tools'])}`")
            if skills.get("cloud"):
                st.markdown(f"**Cloud Services:** `{'`, `'.join(skills['cloud'])}`")
            if skills.get("other"):
                st.markdown(f"**Architecture:** `{'`, `'.join(skills['other'])}`")
        st.divider()

    # Professional Summary
    summary = resume_data.get("professional_summary", "")
    if summary:
        st.markdown("### 📝 Profile Summary")
        st.write(summary)
        st.divider()

    # Featured Technical Projects
    projects = resume_data.get("projects", [])
    if projects:
        st.markdown("### 🛠️ Featured Technical Repositories & Projects")
        for proj in projects:
            p_name = proj.get("project_name", "")
            tech_stack = proj.get("tech_stack", [])
            desc = proj.get("description", "")

            st.markdown(f"#### 📦 {p_name}")
            if tech_stack:
                st.markdown(f"**Tech Stack:** `{'`, `'.join(tech_stack)}`")
            st.write(desc)
            st.write("")
        st.divider()

    # Engineering Experience
    experiences = resume_data.get("experience", [])
    if experiences:
        st.markdown("### 🚀 Engineering Experience")
        for exp in experiences:
            company = exp.get("company", "")
            role = exp.get("role", "")
            start = exp.get("start_date", "")
            end = exp.get("end_date", "")
            loc = exp.get("location", "")

            st.markdown(f"#### {role} @ {company}")
            st.caption(f"{loc} | {start} – {end}")

            for b in exp.get("bullets", []):
                st.markdown(f"- {b}")
            st.write("")
        st.divider()

    # Education Section
    education = resume_data.get("education", [])
    if education:
        st.markdown("### 🎓 Computer Science Education")
        for edu in education:
            degree = edu.get("degree", "")
            inst = edu.get("institution", "")
            grad = edu.get("graduation_date", "")
            gpa = edu.get("gpa", "")

            st.markdown(f"**{degree}** — {inst}")
            st.caption(f"Graduated: {grad} | GPA: {gpa}")


def render_pdf(resume_data: Dict[str, Any], styles: Any, story: list) -> None:
    """Appends Developer Specialist tech-focused ReportLab flowables to story.
    Applies html.escape() to protect against XML special character crashes.

    Args:
        resume_data: Candidate optimized_resume structure.
        styles: ReportLab sample stylesheet.
        story: List of flowable elements to append to.
    """
    personal = resume_data.get("personal_information", {})
    c_name = html.escape(personal.get("name", "Candidate Name"))
    headline = html.escape(resume_data.get("headline", "Full Stack AI Engineer"))
    email = html.escape(personal.get("email", ""))
    phone = html.escape(personal.get("phone", ""))
    location = html.escape(personal.get("location", ""))
    linkedin = html.escape(personal.get("linkedin", ""))
    github = html.escape(personal.get("github", ""))
    portfolio = html.escape(personal.get("portfolio", ""))

    sky_blue = colors.HexColor("#0284C7")

    title_style = ParagraphStyle("Dev_Title", parent=styles["Heading1"], fontName="Helvetica-Bold", fontSize=20, leading=24, textColor=sky_blue, spaceAfter=2)
    sub_style = ParagraphStyle("Dev_Sub", parent=styles["Normal"], fontName="Helvetica-Bold", fontSize=11, leading=14, textColor=colors.HexColor("#334155"), spaceAfter=3)
    handle_style = ParagraphStyle("Dev_Handle", parent=styles["Normal"], fontName="Courier-Bold", fontSize=9, leading=12, textColor=sky_blue, spaceAfter=6)
    section_h2 = ParagraphStyle("Dev_H2", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=11, leading=14, textColor=sky_blue, spaceBefore=8, spaceAfter=3)
    body_style = ParagraphStyle("Dev_Body", parent=styles["Normal"], fontName="Helvetica", fontSize=9, leading=12, textColor=colors.HexColor("#1E293B"), spaceAfter=3)
    bullet_style = ParagraphStyle("Dev_Bullet", parent=styles["Normal"], fontName="Helvetica", fontSize=8.5, leading=11, textColor=colors.HexColor("#334155"), leftIndent=10, spaceAfter=2)

    # Header
    story.append(Paragraph(f"CODE // {c_name}", title_style))
    story.append(Paragraph(headline, sub_style))
    story.append(Paragraph(f"GitHub: {github}  •  Portfolio: {portfolio}  •  LinkedIn: {linkedin}", handle_style))
    story.append(Paragraph(f"{location} | {email} | {phone}", body_style))
    story.append(HRFlowable(width="100%", thickness=1.5, color=sky_blue, spaceAfter=8))

    # Tech Matrix at Top
    skills = resume_data.get("skills", {})
    if skills:
        story.append(Paragraph("DEVELOPER TECH MATRIX", section_h2))
        for cat, items in skills.items():
            if items:
                escaped_cat = html.escape(cat.upper())
                escaped_items = ", ".join([html.escape(i) for i in items])
                story.append(Paragraph(f"<b>[{escaped_cat}]:</b> {escaped_items}", body_style))
        story.append(Spacer(1, 4))

    # Summary
    summary = resume_data.get("professional_summary", "")
    if summary:
        story.append(Paragraph("PROFILE SUMMARY", section_h2))
        story.append(Paragraph(html.escape(summary), body_style))
        story.append(Spacer(1, 4))

    # Featured Projects
    projects = resume_data.get("projects", [])
    if projects:
        story.append(Paragraph("FEATURED TECHNICAL REPOSITORIES & PROJECTS", section_h2))
        for proj in projects:
            p_name = html.escape(proj.get("project_name", ""))
            tech_stack = ", ".join([html.escape(t) for t in proj.get("tech_stack", [])])
            desc = html.escape(proj.get("description", ""))

            story.append(Paragraph(f"<b>REPOS // {p_name}</b> <i>[{tech_stack}]</i>", body_style))
            story.append(Paragraph(desc, bullet_style))
            story.append(Spacer(1, 3))

    # Engineering Experience
    experiences = resume_data.get("experience", [])
    if experiences:
        story.append(Paragraph("ENGINEERING EXPERIENCE", section_h2))
        for exp in experiences:
            comp = html.escape(exp.get("company", ""))
            role = html.escape(exp.get("role", ""))
            start = html.escape(exp.get("start_date", ""))
            end = html.escape(exp.get("end_date", ""))
            loc = html.escape(exp.get("location", ""))

            story.append(Paragraph(f"<b>{role} @ {comp}</b> (<i>{loc} | {start} – {end}</i>)", body_style))
            for b in exp.get("bullets", []):
                story.append(Paragraph(f"• {html.escape(b)}", bullet_style))
            story.append(Spacer(1, 3))

    # Education
    education = resume_data.get("education", [])
    if education:
        story.append(Paragraph("COMPUTER SCIENCE EDUCATION", section_h2))
        for edu in education:
            deg = html.escape(edu.get("degree", ""))
            inst = html.escape(edu.get("institution", ""))
            grad = html.escape(edu.get("graduation_date", ""))
            gpa = html.escape(str(edu.get("gpa", "")))

            story.append(Paragraph(f"<b>{deg}</b> — {inst} ({grad}) {f'| GPA: {gpa}' if gpa else ''}", body_style))


def render_docx(resume_data: Dict[str, Any], doc: Any) -> None:
    """Appends Developer Specialist tech-focused layout to Python-Docx Document.

    Args:
        resume_data: Candidate optimized_resume structure.
        doc: Python-Docx Document instance.
    """
    personal = resume_data.get("personal_information", {})
    c_name = personal.get("name", "Candidate Name")
    headline = resume_data.get("headline", "Full Stack AI Engineer")
    email = personal.get("email", "")
    phone = personal.get("phone", "")
    location = personal.get("location", "")
    linkedin = personal.get("linkedin", "")
    github = personal.get("github", "")
    portfolio = personal.get("portfolio", "")

    sky_blue = RGBColor(2, 132, 199)

    h_name = doc.add_heading(f"CODE // {c_name}", level=0)
    h_name.alignment = WD_ALIGN_PARAGRAPH.LEFT

    p_head = doc.add_paragraph()
    r_head = p_head.add_run(headline)
    r_head.bold = True
    r_head.font.color.rgb = sky_blue
    r_head.font.size = Pt(11)

    p_handles = doc.add_paragraph(f"GitHub: {github} | Portfolio: {portfolio} | LinkedIn: {linkedin}")
    p_handles.style.font.size = Pt(9)
    p_handles.style.font.color.rgb = sky_blue

    doc.add_paragraph(f"{location} | {email} | {phone}")

    # Tech Matrix at Top
    skills = resume_data.get("skills", {})
    if skills:
        doc.add_heading("DEVELOPER TECH MATRIX", level=2)
        for cat, items in skills.items():
            if items:
                p_sk = doc.add_paragraph()
                r_cat = p_sk.add_run(f"[{cat.upper()}]: ")
                r_cat.bold = True
                p_sk.add_run(", ".join(items))

    # Profile Summary
    summary = resume_data.get("professional_summary", "")
    if summary:
        doc.add_heading("PROFILE SUMMARY", level=2)
        doc.add_paragraph(summary)

    # Featured Projects
    projects = resume_data.get("projects", [])
    if projects:
        doc.add_heading("FEATURED TECHNICAL REPOSITORIES & PROJECTS", level=2)
        for proj in projects:
            p_name = proj.get("project_name", "")
            tech_stack = ", ".join(proj.get("tech_stack", []))
            desc = proj.get("description", "")

            p_proj = doc.add_paragraph()
            p_proj.add_run(f"REPOS // {p_name}").bold = True
            if tech_stack:
                p_proj.add_run(f" [{tech_stack}]").italic = True
            doc.add_paragraph(desc)

    # Engineering Experience
    experiences = resume_data.get("experience", [])
    if experiences:
        doc.add_heading("ENGINEERING EXPERIENCE", level=2)
        for exp in experiences:
            comp = exp.get("company", "")
            role = exp.get("role", "")
            start = exp.get("start_date", "")
            end = exp.get("end_date", "")
            loc = exp.get("location", "")

            p_exp = doc.add_paragraph()
            r_role = p_exp.add_run(f"{role} @ {comp}")
            r_role.bold = True
            p_exp.add_run(f" ({loc} | {start} – {end})").italic = True

            for b in exp.get("bullets", []):
                doc.add_paragraph(b, style="List Bullet")

    # Education
    education = resume_data.get("education", [])
    if education:
        doc.add_heading("COMPUTER SCIENCE EDUCATION", level=2)
        for edu in education:
            deg = edu.get("degree", "")
            inst = edu.get("institution", "")
            grad = edu.get("graduation_date", "")
            gpa = edu.get("gpa", "")

            p_edu = doc.add_paragraph()
            p_edu.add_run(f"{deg} — {inst}").bold = True
            p_edu.add_run(f" ({grad}) {f'GPA: {gpa}' if gpa else ''}")
